import os
#!/usr/bin/env python3
# Compares Survey123 .xlsx names with existing Word feature service templates
# and writes a human-readable TXT report named by concatenating the input filenames.

from pathlib import Path
import re
import sys
import unicodedata
from openpyxl import load_workbook
from docx import Document

# ──────────────────────────────────────────────────────────────────────────────
# CONFIGURATION (edit these values to suit your environment)
# ──────────────────────────────────────────────────────────────────────────────

# Input .xlsx and .docx paths (hard-coded variables)
XLSX_PATH = Path(r"C:\Users\Michael.Fleming\Aurecon Group\523617 - M12RT IC - Black Hill to Tomago - GIS\Working\MF\FRT\PRJ_523617_M12RT_Black_Hill_to _Tomago.xlsx")
DOCX_PATH = Path(r"C:\Users\Michael.Fleming\Aurecon Group\523617 - M12RT IC - Black Hill to Tomago - GIS\Working\MF\FRT\M12RT_BH2T_template_01.docx")

# Comparison options
IGNORE_CASE = True          # lowercases both sides before comparing
NORMALIZE_SPACES = True     # replace spaces with underscores before comparing

# Where to place the TXT report (by default, next to the XLSX file).
# The filename is "<xlsx_stem>__<docx_stem>.txt" with disallowed characters removed.
OUTPUT_DIR = XLSX_PATH.parent  # change if you prefer another folder

# ──────────────────────────────────────────────────────────────────────────────
# IMPLEMENTATION
# ──────────────────────────────────────────────────────────────────────────────

def read_survey_names_from_xlsx(xlsx_path: Path):
    """
    Read column B (index 2) from the sheet named 'survey' (case-insensitive).
    Start at row 2 (skip header). Preserve order and dedupe (first occurrence wins).
    Returns list[str].
    """
    from openpyxl import load_workbook

    wb = load_workbook(filename=str(xlsx_path), read_only=True, data_only=True)
    survey_sheet = None
    for name in wb.sheetnames:
        if name.lower() == "survey":
            survey_sheet = wb[name]
            break
    if survey_sheet is None:
        raise ValueError(f"No sheet named 'survey' found in {xlsx_path}. Available sheets: {wb.sheetnames}")
    names = []
    seen = set()
    for row in survey_sheet.iter_rows(min_row=2, min_col=2, max_col=2, values_only=True):
        val = row[0]
        if val is None:
            continue
        s = str(val).strip()
        if s == "":
            continue
        if s not in seen:
            seen.add(s)
            names.append(s)
    return names

def _strip_combining_marks(s: str) -> str:
    """Normalize to NFD and remove combining marks so diacritics are collapsed for matching."""
    n = unicodedata.normalize("NFD", s)
    return "".join(ch for ch in n if not unicodedata.combining(ch))

def _extract_field_from_token_content(content: str):
    """
    Given the text between ${ and }, extract the field name according to rules:
    - Skip optional leading whitespace.
    - If content starts with '#' -> discard it and leading whitespace.
    - If content starts with 'if' (case-insensitive; diacritics handled), discard it.
    - Field name is the first whole word [a-z_]+.
    Returns the field name (lowercase) or None.
    """
    if content is None:
        return None

    content_original = content
    content_norm = _strip_combining_marks(content_original)
    content_norm_stripped = content_norm.lstrip()

    # Remove leading '#'
    if content_norm_stripped.startswith("#"):
        orig_stripped = content_original.lstrip()
        orig_after_hash = orig_stripped[1:].lstrip() if orig_stripped.startswith("#") else orig_stripped
        search_source = orig_after_hash
    else:
        low = content_norm_stripped.lower()
        if low.startswith("if") and (len(low) == 2 or (len(low) > 2 and not low[2].isalpha())):
            orig_stripped = content_original.lstrip()
            m = re.match(r'(?i)if\b', _strip_combining_marks(orig_stripped))
            if m:
                cut_len = len(m.group(0))
                orig_after_if = orig_stripped[cut_len:].lstrip()
                search_source = orig_after_if
            else:
                search_source = content_original.lstrip()
        else:
            search_source = content_original.lstrip()

    search_norm = _strip_combining_marks(search_source).lower()
    m2 = re.search(r'\b([a-z_]+)\b', search_norm)
    if not m2:
        return None
    candidate = m2.group(1)
    return candidate

def extract_field_tokens_from_docx(docx_path: Path):
    """
    Scan paragraphs and table cells for ${...} occurrences.
    Return list of unique field names preserving first-seen order.
    """
    from docx import Document

    doc = Document(str(docx_path))
    token_pattern = re.compile(r'\$\{\s*([^\}]*)\s*\}')
    found = []
    seen = set()

    def scan_text_block(text: str):
        for m in token_pattern.finditer(text):
            inner = m.group(1)
            fld = _extract_field_from_token_content(inner)
            if fld and fld not in seen:
                seen.add(fld)
                found.append(fld)

    for p in doc.paragraphs:
        if p.text:
            scan_text_block(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    scan_text_block(cell.text)

    return found

def normalize_list(items, ignore_case=False, normalize_spaces=False):
    out = []
    seen = set()
    for it in items:
        s = it
        if s is None:
            continue
        s = s.strip()
        if normalize_spaces:
            s = s.replace(" ", "_")
        if ignore_case:
            s = s.lower()
        if s and s not in seen:
            seen.add(s)
            out.append(s)
    return out

def compare_lists(xlsx_list, docx_list):
    set_x = set(xlsx_list)
    set_d = set(docx_list)
    in_both = [s for s in xlsx_list if s in set_d]  # preserve XLSX order
    only_x = sorted(list(set_x - set_d))
    only_d = sorted(list(set_d - set_x))
    return {
        "xlsx_total": len(xlsx_list),
        "docx_total": len(docx_list),
        "in_both": in_both,
        "only_in_xlsx": only_x,
        "only_in_docx": only_d,
    }

# ──────────────────────────────────────────────────────────────────────────────
# TXT report helpers
# ──────────────────────────────────────────────────────────────────────────────

def _safe_stem(p: Path) -> str:
    """
    Return a filesystem-safe stem for a path:
    - use Path.stem
    - replace characters not in [A-Za-z0-9._-] with '_'
    """
    stem = p.stem
    return re.sub(r'[^A-Za-z0-9._-]+', '_', stem)

def build_output_txt_path(xlsx_path: Path, docx_path: Path, out_dir: Path) -> Path:
    """
    Concatenate stems to form: <xlsx_stem>__<docx_stem>.txt
    """
    xstem = _safe_stem(xlsx_path)
    dstem = _safe_stem(docx_path)
    fname = f"{xstem}__{dstem}.txt"
    return out_dir / fname

def write_text_report(path: Path, report: dict, xlsx_raw, docx_raw, xlsx_path: Path, docx_path: Path,
                      ignore_case: bool, normalize_spaces: bool):
    """
    Write a readable text file with summary and differences.
    """
    lines = []
    lines.append("Survey123 vs Template Field Comparison")
    lines.append("=" * 42)
    lines.append("")
    lines.append(f"XLSX file: {xlsx_path}")
    lines.append(f"DOCX file: {docx_path}")
    lines.append(f"Options: ignore_case={ignore_case}, normalize_spaces={normalize_spaces}")
    lines.append("")
    lines.append("Summary")
    lines.append("-------")
    lines.append(f"Total fields in XLSX (survey col B): {report['xlsx_total']}")
    lines.append(f"Total fields in DOCX (${ { '{...}' } } tokens): {report['docx_total']}")
    lines.append("")
    lines.append(f"In both ({len(report['in_both'])}):")
    for v in report["in_both"]:
        lines.append(f"  - {v}")
    lines.append("")
    lines.append(f"Only in XLSX ({len(report['only_in_xlsx'])}):")
    for v in report["only_in_xlsx"]:
        lines.append(f"  - {v}")
    lines.append("")
    lines.append(f"Only in DOCX ({len(report['only_in_docx'])}):")
    for v in report["only_in_docx"]:
        lines.append(f"  - {v}")
    lines.append("")
    lines.append("Raw extracted values (for audit)")
    lines.append("-------------------------------")
    lines.append("XLSX raw:")
    for v in xlsx_raw:
        lines.append(f"  - {v}")
    lines.append("DOCX raw:")
    for v in docx_raw:
        lines.append(f"  - {v}")

    text = "\n".join(lines)
    path.write_text(text, encoding="utf-8")

# ──────────────────────────────────────────────────────────────────────────────

def main():
    # Validate files exist
    if not XLSX_PATH.exists():
        print(f"Error: XLSX not found: {XLSX_PATH}", file=sys.stderr)
        return 2
    if not DOCX_PATH.exists():
        print(f"Error: DOCX not found: {DOCX_PATH}", file=sys.stderr)
        return 2

    # Extract & normalize
    try:
        xlsx_raw = read_survey_names_from_xlsx(XLSX_PATH)
    except Exception as e:
        print(f"Error reading xlsx: {e}", file=sys.stderr)
        return 3

    try:
        docx_raw = extract_field_tokens_from_docx(DOCX_PATH)
    except Exception as e:
        print(f"Error reading docx: {e}", file=sys.stderr)
        return 4

    xlsx_list = normalize_list(xlsx_raw, ignore_case=IGNORE_CASE, normalize_spaces=NORMALIZE_SPACES)
    docx_list = normalize_list(docx_raw, ignore_case=IGNORE_CASE, normalize_spaces=NORMALIZE_SPACES)

    report = compare_lists(xlsx_list, docx_list)

    # Build output text file path by concatenating input base names
    OUTPUT_TXT = build_output_txt_path(XLSX_PATH, DOCX_PATH, OUTPUT_DIR)

    # Print summary to console
    print("Comparison summary")
    print("------------------")
    print(f"Fields extracted from XLSX (survey sheet, column B): {report['xlsx_total']}")
    print(f"Fields extracted from DOCX (${ { '{...}' } } tokens): {report['docx_total']}")
    print()
    print(f"In both ({len(report['in_both'])}):")
    for v in report["in_both"]:
        print(f"  - {v}")
    print()
    print(f"Only in XLSX ({len(report['only_in_xlsx'])}):")
    for v in report["only_in_xlsx"]:
        print(f"  - {v}")
    print()
    print(f"Only in DOCX ({len(report['only_in_docx'])}):")
    for v in report["only_in_docx"]:
        print(f"  - {v}")

    # Write text report
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        write_text_report(
            OUTPUT_TXT, report, xlsx_raw, docx_raw,
            XLSX_PATH, DOCX_PATH, IGNORE_CASE, NORMALIZE_SPACES
        )
        print()
        txt_path = OUTPUT_TXT.resolve()
        print(f"Text report written to: {txt_path}")
        os.startfile(txt_path)
    except Exception as e:
        print(f"Error writing text report: {e}", file=sys.stderr)
        return 5

    return 0

if __name__ == "__main__":
    sys.exit(main() or 0)
